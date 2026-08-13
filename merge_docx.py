#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
合并润色稿 → 分类 docx（宋体正文 + 黑体标题，用户中文 Word 铁律）。
分类规则：按标题关键词分 6 类；每类内按原序号排序。
用法: python merge_docx.py [--out 输出路径] [--src 润色稿目录]
"""
import argparse
import re
import sys
from pathlib import Path

# 分类：关键词 → 类别名（按匹配顺序，先匹配先得）
CATEGORIES = [
    ("盾构与TBM掘进", ["盾构", "TBM", "掘进", "刀盘", "顶管", "联络通道", "推力壳", "竖井掘进", "沉井", "垂直掘进"]),
    ("隧道施工工法", ["冻结", "降水", "明挖", "暗挖", "矿山法", "管片", "装配式", "连续墙", "机械化", "注浆", "堵漏", "防水", "渗漏"]),
    ("岩土与基坑工程", ["岩土", "基坑", "软土", "地基", "地下连续墙", "沉降", "变形", "监测", "传感"]),
    ("地下空间与储能", ["地下空间", "储能", "压气", "地下洞室", "车站", "地铁"]),
    ("工程管理与数字化", ["数字化", "智能", "AI", "人工智能", "信息化", "管理", "风险", "保险", "市场", "产业", "转型", "国际化"]),
    ("人物访谈与行业见闻", ["对话", "访谈", "故事", "人生", "经历", "行业", "经验", "职业", "风采", "游历", "见闻"]),
]
DEFAULT_CAT = "其他"


def classify(title: str) -> str:
    for name, kws in CATEGORIES:
        if any(k in title for k in kws):
            return name
    return DEFAULT_CAT


def sort_key(path: Path):
    m = re.match(r"(\d{3})_(.+)", path.stem)
    if m:
        return (int(m.group(1)), path.stem)
    return (999, path.stem)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--src", default=r"D:\Documents\hermes_work\YWork\videosays-local\outputs\polished")
    ap.add_argument("--out", default=r"D:\Documents\hermes_work\YWork\videosays-local\outputs\白云说隧_讲稿合集_润色版.docx")
    args = ap.parse_args()

    src = Path(args.src)
    files = sorted([f for f in src.glob("*.txt") if not f.name.startswith("group")],
                   key=sort_key)
    if not files:
        print("未找到润色稿")
        sys.exit(1)

    # 分类
    buckets = {}
    for f in files:
        title = f.stem
        cat = classify(title)
        buckets.setdefault(cat, []).append(f)
    print("分类统计:")
    for cat, fs in buckets.items():
        print(f"  {cat}: {len(fs)} 篇")

    # 生成 docx
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    doc = Document()
    # 全局正文样式：宋体
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 文档总标题（黑体）
    h = doc.add_paragraph()
    r = h.add_run("白云说隧·隧道访谈讲稿合集（润色版）")
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(20)
    r.font.bold = True

    for cat in buckets:
        # 类别标题（黑体 16）
        h2 = doc.add_paragraph()
        r2 = h2.add_run(cat)
        r2.font.name = "黑体"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        r2.font.size = Pt(16)
        r2.font.bold = True

        for f in sorted(buckets[cat], key=sort_key):
            # 篇名（黑体 14）
            h3 = doc.add_paragraph()
            r3 = h3.add_run(f.stem)
            r3.font.name = "黑体"
            r3._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            r3.font.size = Pt(14)
            r3.font.bold = True

            # 正文（宋体 12）
            text = f.read_text(encoding="utf-8")
            for para in text.split("\n"):
                para = para.strip()
                if not para:
                    continue
                p = doc.add_paragraph()
                run = p.add_run(para)
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(12)

    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"\n✅ 已输出: {out} ({out.stat().st_size/1024:.0f} KB, {len(files)} 篇)")


if __name__ == "__main__":
    main()
