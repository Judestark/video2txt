#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""videosays_local.asr — faster-whisper 本地转写核心
输出 txt / srt / vtt 三种格式，与 videosays 云端一致。
运行方式: python -m videosays_local.asr --wav <file> --model small --out <txt>
"""
import argparse
import os
import sys
import time
from pathlib import Path


def fmt_ts(seconds: float, sep: str = ",") -> str:
    ms = int(round((seconds - int(seconds)) * 1000))
    s = int(seconds) % 60
    m = (int(seconds) // 60) % 60
    h = int(seconds) // 3600
    return f"{h:02d}:{m:02d}:{s:02d}{sep}{ms:03d}"


def write_docx(out_base: str, title: str, segments: list):
    """生成 Word 文档：黑体标题 + 宋体正文（用户中文 Word 铁律）。
    segments: [(text, start, end), ...]
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
    except ImportError:
        print("[docx] python-docx 未安装，跳过 .docx 输出", flush=True)
        return

    doc = Document()
    # 页面默认样式：正文宋体
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 标题：黑体
    h = doc.add_paragraph()
    hr = h.add_run(title)
    hr.font.name = "黑体"
    hr._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    hr.font.size = Pt(16)
    hr.font.bold = True

    for text, start, end in segments:
        if not text.strip():
            continue
        p = doc.add_paragraph()
        # 时间戳段（灰色小字）+ 正文
        ts = p.add_run(f"[{fmt_ts(start)} - {fmt_ts(end)}] ")
        ts.font.size = Pt(9)
        ts.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        body = p.add_run(text)
        body.font.name = "Times New Roman"
        body._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        body.font.size = Pt(12)

    doc.save(out_base + ".docx")


def srt_block(i: int, start: float, end: float, text: str) -> str:
    return f"{i}\n{fmt_ts(start)} --> {fmt_ts(end)}\n{text}\n"


def vtt_block(i: int, start: float, end: float, text: str) -> str:
    return f"{fmt_ts(start, '.')} --> {fmt_ts(end, '.')}\n{text}\n"


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--wav", required=True, help="输入 16k wav")
    ap.add_argument("--model", default="small", help="small/medium/large-v3")
    ap.add_argument("--out", required=True, help="输出 .txt 路径（自动生成 .srt/.vtt）")
    ap.add_argument("--language", default="zh", help="zh/en/auto")
    ap.add_argument("--device", default="auto", help="auto/cpu/cuda")
    ap.add_argument("--compute-type", default="auto", help="auto/int8/float16")
    ap.add_argument("--prompt", default="", help="领域术语提示词（提高术语准确率）")
    args = ap.parse_args()

    # 模型缓存目录固定，避免每次询问；国内走 hf-mirror 镜像
    cache_dir = os.environ.get("HF_HOME") or os.path.join(
        os.path.expanduser("~"), ".cache", "huggingface")
    os.environ["HF_HOME"] = cache_dir
    os.environ.setdefault("HF_ENDPOINT", "https://hf-mirror.com")
    os.environ.setdefault("HF_HUB_DISABLE_PROGRESS_BARS", "1")

    # device 自动选择：有 GPU 用 cuda
    if args.device == "auto":
        try:
            import ctranslate2
            args.device = "cuda" if ctranslate2.get_cuda_device_count() > 0 else "cpu"
        except Exception:
            args.device = "cpu"
    # 模型解析：优先用本地 models/ 目录（ModelScope 预下载），否则 HF 下载
    model_ref = args.model
    local_dir = Path(__file__).resolve().parent.parent / "models" / f"faster-whisper-{args.model}"
    if local_dir.exists() and (local_dir / "model.bin").exists():
        model_ref = str(local_dir)
        print(f"[asr] 使用本地模型: {model_ref}")

    t0 = time.time()
    print(f"[asr] 加载模型 {args.model} (device={args.device}, 首次会自动下载)", flush=True)
    from faster_whisper import WhisperModel
    model = WhisperModel(model_ref, device=args.device, compute_type=args.compute_type)
    print(f"[asr] 模型就绪 {time.time()-t0:.0f}s，开始转写...", flush=True)

    prompt = args.prompt or ("以下是普通话的讲座、旁白或访谈内容，涉及隧道工程、土木工程、建筑施工、"
                             "基坑工程等专业领域。常见术语：盾构法、矿山法、明挖法、暗挖法、深基坑、"
                             "基坑围护、钢筋混凝土、钢结构、型钢、混凝土、围护结构、接缝渗漏、"
                             "地下水位、挡土、抗震、延性、地域性、制约、他山之石可以攻玉、"
                             "施工材料、隧道建造技术、闲余时间。")
    segments, info = model.transcribe(args.wav, language=args.language,
                                      vad_filter=True,
                                      beam_size=5,
                                      initial_prompt=prompt)
    segs = list(segments)

    # 后处理：去除相邻重复段（VAD 边界重叠导致的整段重复）
    dedup = []
    for s in segs:
        t = s.text.strip()
        if dedup and t == dedup[-1][0]:
            continue  # 与前一段完全相同 -> 跳过
        dedup.append((t, s.start, s.end))

    # 合并为三段输出
    txt = "\n".join(t for t, _, _ in dedup)
    srt = "\n".join(srt_block(i, st, en, t).rstrip()
                    for i, (t, st, en) in enumerate(dedup, 1)) + "\n"
    vtt = "WEBVTT\n\n" + "\n".join(vtt_block(i, st, en, t).rstrip()
                                   for i, (t, st, en) in enumerate(dedup, 1)) + "\n"

    base = os.path.splitext(args.out)[0]
    with open(args.out, "w", encoding="utf-8") as f:
        f.write(txt)
    with open(base + ".srt", "w", encoding="utf-8") as f:
        f.write(srt)
    with open(base + ".vtt", "w", encoding="utf-8") as f:
        f.write(vtt)

    # Word 输出：黑体标题 + 宋体正文（时间戳灰色）
    title = os.path.basename(base)
    write_docx(base, title, dedup)

    lang = getattr(info, "language", "?")
    print(f"[asr] ✅ {time.time()-t0:.0f}s 完成，语言={lang}，"
          f"{len(segs)} 段 -> {args.out} / .srt / .vtt / .docx", flush=True)


if __name__ == "__main__":
    main()
